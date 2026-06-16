"""
rag.py - RAG 知识库模块
使用 LangChain + LanceDB 向量库 + SiliconFlow Embedding API
"""
import os
import io
import uuid
from typing import List

LANCE_PATH = "./lance_db"
TABLE_NAME = "astro_knowledge"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

_vectorstore = None
_embeddings = None


def _get_embeddings():
    global _embeddings
    if _embeddings is not None:
        return _embeddings
    import httpx
    from langchain_openai import OpenAIEmbeddings
    _embeddings = OpenAIEmbeddings(
        model="BAAI/bge-m3",
        openai_api_key=os.getenv("AI_API_KEY", ""),
        openai_api_base=os.getenv("AI_BASE_URL", "https://api.siliconflow.cn/v1"),
        tiktoken_enabled=False,
        check_embedding_ctx_length=False,
        http_client=httpx.Client(trust_env=False),
    )
    return _embeddings


def _get_vectorstore():
    global _vectorstore
    if _vectorstore is not None:
        return _vectorstore

    import lancedb
    from langchain_community.vectorstores import LanceDB

    embeddings = _get_embeddings()
    db = lancedb.connect(LANCE_PATH)

    if TABLE_NAME in db.table_names():
        _vectorstore = LanceDB(
            connection=db,
            embedding=embeddings,
            table_name=TABLE_NAME,
            vector_key="vector",
            id_key="id",
            text_key="text",
            mode="append",
        )
    else:
        from langchain_core.documents import Document
        # LanceDB 必须用真实文档建表，无法空建；写入一条占位符后立即删除
        placeholder = Document(
            page_content="__init__",
            metadata={"doc_id": "__init__", "filename": "", "chunk_index": 0},
        )
        _vectorstore = LanceDB.from_documents(
            documents=[placeholder],
            embedding=embeddings,
            connection=db,
            table_name=TABLE_NAME,
            vector_key="vector",
            id_key="id",
            text_key="text",
            mode="append",
        )
        db.open_table(TABLE_NAME).delete("metadata.doc_id = '__init__'")

    return _vectorstore


def _chunk_text(text: str) -> List[str]:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
    )
    return [c for c in splitter.split_text(text) if c.strip()]


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")

    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if ext in ("docx", "doc"):
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)

    raise ValueError(f"不支持的文件类型: {ext}")


def add_document(doc_id: str, text: str, metadata: dict) -> int:
    from langchain_core.documents import Document

    chunks = _chunk_text(text)
    if not chunks:
        return 0

    # 先删旧数据，支持同一文档重复上传时的幂等更新
    delete_document(doc_id)

    filename = metadata.get("filename", "")
    docs = [
        Document(
            page_content=chunk,
            metadata={
                "doc_id": doc_id,
                "filename": filename,
                "chunk_index": i,
            },
        )
        for i, chunk in enumerate(chunks)
    ]

    vs = _get_vectorstore()
    ids = [f"{doc_id}_chunk_{i}" for i in range(len(docs))]
    vs.add_documents(docs, ids=ids)

    return len(chunks)


def delete_document(doc_id: str):
    try:
        import lancedb
        db = lancedb.connect(LANCE_PATH)
        if TABLE_NAME not in db.table_names():
            return
        table = db.open_table(TABLE_NAME)
        table.delete(f"metadata.doc_id = '{doc_id}'")
    except Exception:
        pass


def get_chunks(doc_id: str) -> List[dict]:
    """返回某文档所有切块，按 chunk_index 排序"""
    try:
        import lancedb
        db = lancedb.connect(LANCE_PATH)
        if TABLE_NAME not in db.table_names():
            return []
        table = db.open_table(TABLE_NAME)
        df = table.to_pandas()
        rows = df[df["metadata"].apply(lambda m: m.get("doc_id") == doc_id)]
        return sorted(
            [{"chunk_index": r["metadata"]["chunk_index"], "text": r["text"]} for _, r in rows.iterrows()],
            key=lambda x: x["chunk_index"],
        )
    except Exception:
        return []


def search(query: str, n_results: int = 3) -> List[str]:
    try:
        import lancedb
        db = lancedb.connect(LANCE_PATH)
        if TABLE_NAME not in db.table_names():
            return []
        table = db.open_table(TABLE_NAME)
        if table.count_rows() == 0:
            return []
    except Exception:
        return []

    vs = _get_vectorstore()
    docs = vs.similarity_search(query, k=n_results)
    return [doc.page_content for doc in docs]
